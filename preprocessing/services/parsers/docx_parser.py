import io
from datasets import Dataset
from .base import BaseParser
import docx

class DOCXParser(BaseParser):
    """Parser for Microsoft Word (DOCX) documents."""

    def parse(self, file_stream: io.BytesIO) -> Dataset:
        """Parse a DOCX file from an in-memory stream.

        Args:
            file_stream: An in-memory binary stream (e.g., io.BytesIO)
                containing the DOCX file content.

        Returns:
            Dataset: with a 'text' column.
        """
        try:
            doc = docx.Document(file_stream)
            data = []
            
            for p in doc.paragraphs:
                if p.text:
                    data.append({"text": p.text.strip()})
            
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            data.append({"text": cell.text.strip()})

            return Dataset.from_list(data)

        except Exception as e:
            print(f"Error parsing DOCX file {file_path}: {e}")
            return Dataset.from_list([])